import docx
import os
import tkinter as tk
from tkinter import *
import ttkbootstrap as ttk
from ttkbootstrap.constants import *

import tkinter.messagebox as mb

from tkinter import filedialog as fd

from PyPDF2 import PdfReader
import pyexcel_xls
import pyexcel

from win32com import client as wc

import threading

from .config import logger


def callback():
    return fd.askdirectory()


def find_all(chose_dir):

    paths = []
    paths_doc = []
    paths_pdf = []
    paths_xls = [] 

    for root, dirs, files in chose_dir:

        for file in files:
            if file.endswith(('docx')) and not file.startswith('~'):
                paths.append(os.path.join(root, file))
            if file.endswith(('doc')) and not file.startswith('~'):
                paths_doc.append(os.path.join(root, file))
            if file.endswith('pdf') and not file.startswith('~'):
                paths_pdf.append(os.path.join(root, file))
            if file.endswith(('xls', 'xlsx')) and not file.startswith('~'):
                paths_xls.append(os.path.join(root, file))
    
    return paths, paths_pdf, paths_xls, paths_doc



# def find_only_dir(chosen_dir):
    
#     paths = []
#     paths_pdf = []
#     paths_xls = []

#     # Обходим все файлы в выбранной директории, но не во вложенных папках
#     for root, dirs, files in chosen_dir:

#         for file in files:
#             if file.endswith('docx') and not file.startswith('~'):
#                 paths.append(os.path.join(root, file))
#             if file.endswith('pdf') and not file.startswith('~'):
#                 paths_pdf.append(os.path.join(root, file))
#             if file.endswith(('xls', 'xlsx')) and not file.startswith('~'):
#                 paths_xls.append(os.path.join(root, file))
#         break
    
#     return paths, paths_pdf, paths_xls



def find_docx(find, txt1, path):

    logger.info('Check .docx')
    try:
        doc = docx.Document(path)
        for paragraph in doc.paragraphs:
            
            paragraph = paragraph.text.lower().replace('\n', '').replace(' ', '')

            logger.info(f'{find=}')
            logger.info(f'{paragraph.text.lower()=}')

            if find in paragraph:

                not_found = 0
                flag = 1
                path, filename = os.path.split(path)
                update_result_text(filename, path, txt1)
                break

        if flag == 0:  # если не нашли в параграфах, ищем в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:

                        if find in cell.text.lower().replace('\n', '').replace(' ', ''):
                            not_found = 0
                            path, filename = os.path.split(path)
                            update_result_text(filename, path, txt1)
                            break
                    if not_found == 0:
                        break
                if not_found == 0:
                    break
    except Exception as e:
        logger.error(f'Error docx: {e}')

def find_doc(find, txt1, path):
    try:
        logger.info(f'{txt1=}')
        w = wc.Dispatch('Word.Application')

        doc=w.Documents.Open(os.path.abspath(path))

        text = doc.Content.Text.lower().replace('\n', '').replace(' ', '').replace('\r', '')
        logger.info(f'{text=}')
        cell = text.find(find)
        logger.info(f'{cell=}')
        if find != -1:
            logger.info('Нашли')
            path, filename = os.path.split(path)
            update_result_text(filename, path, txt1)
            not_found = 0
        
        if not_found == 0:
            return not_found

    except Exception as e:
        logger.error(f"Error reading or decoding the doc file: {e}")

    finally:
        # Важно закрыть файл, когда работа с ним завершена
        w.Quit()

    
def find_pdf(find, txt1, path):
    logger.info('Check .pdf')
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            clear_page = page.extract_text().replace('\n', '').replace(' ', '')


            logger.info(f'{find=}')
            logger.info(f'{clear_page.lower()=}')

            if find in clear_page.lower():
                not_found = 0
                path, filename = os.path.split(path)
                update_result_text(filename, path, txt1)
                break
        
        if not_found == 0:
            return not_found
    except Exception as e:
        logger.error(f'Error pdf: {e}')
    

def findDocs_process(chose_dir, find, txt1, search_mode):

    find = find.replace('\n', '').replace(' ', '')
    logger.info(f'{find=}')

    flag_xls = 0
    flag = 0
    not_found = 1
    
    logger.info(f'Find_all_function_without_xlsx')
    paths, paths_pdf, paths_xls, paths_doc = find_all(chose_dir)

    logger.info(f'{paths=}\n\n{paths_pdf=}\n\n{paths_xls=}')
    # поиск по docx
    for path in paths:
        result_function = find_docx(find, txt1, path)
        if result_function == 0:
            not_found = 0
    
    for path in paths_doc:
        result_function = find_doc(find, txt1, path)
        if result_function == 0:
            not_found = 0

    # поиск по pdf
    for path in paths_pdf:
        result_function = find_pdf(find, txt1, path)
        if result_function == 0:
            not_found = 0
        

    # поиск по xls
    # for path in paths_xls:
    #     logger.info('Check .xls and .xlsx')
    #     try:
    #         excel_file = pyexcel.get_array(file_name=path)
    #         for elem in excel_file:
    #             for cell in elem:

    #                 try:
    #                     result = chardet.detect(str(cell).encode('utf-8'))
    #                     logger.info(f'{result=}')
    #                     # Преобразование текста для проверки корректности кодировки
    #                     logger.info(f'{str(cell).encode('utf-8').decode('utf-8')=}')
    #                     # logger.info(f'{str(cell).encode('utf-8').decode('utf-8')=}')
    #                 except UnicodeDecodeError:
    #                     logger.error(f"File '{path}' contains non-UTF-8 characters.")
    #                     logger.error(f'{str(cell)[:100]=}')

    #                 if find in str(cell).lower():
    #                     not_found = 0
    #                     path, filename = os.path.split(path)
    #                     update_result_text(filename, path, txt1)
    #                     flag_xls = 1
    #                     break
    #             if flag_xls == 1:
    #                 break
    #     except Exception as e:
    #         logger.error(f'Error xls: {e}')

    if not_found == 1:
        mb.showwarning('Ошибка!', 'Не найдено документов, содержащих данное слово/предложение!')
    
    progress_bar.stop()
    progress_bar.pack_forget()

def update_result_text(filename, path, txt1):
    txt1.config(state="normal")
    txt1.insert(tk.END, '\n' + filename + '\n' + path + '\n')
    txt1.config(state="disabled")

def start_search(chose_dir, find, txt1, left_frame1, search_mode):

    global progress_bar
    progress_bar = ttk.Progressbar(left_frame1, bootstyle="success-striped")
    progress_bar.pack(side=TOP, fill=BOTH, padx=10, pady=10)
    progress_bar.start()

    thread = threading.Thread(target=lambda: findDocs_process(chose_dir, find, txt1, search_mode))
    thread.start()

def findDocs(find_entry, txt1, left_frame1, search_mode):

    find = find_entry.get("1.0", tk.END).strip().lower()
    
    if find != '':
        logger.info(f'Start funding: {find=}')
        chose_dir = os.walk(callback())
        start_search(chose_dir, find, txt1, left_frame1, search_mode)
    else:
        mb.showwarning('Ошибка!', 'Вы ввели пустое значение!')